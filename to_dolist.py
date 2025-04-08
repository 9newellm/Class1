from docx import Document

# Create the Word document
doc = Document()
doc.add_heading("Organized To-Do List (April 8–12)", level=1)

# Data structured by day
todo_by_day = {
    "Monday (Today)": [
        "Revise & submit English syllabus (1 hr)",
        "Email professor (15 min)",
        "Figure out the direction of English class (30 min)",
        "Start poster draft using past research (1 hr)",
        "Integrate trauma & Marxism into poster idea (1 hr)",
        "Atomic & Nuclear: Work through current tasks (6 hrs)",
        "Modern Physics: Finalize all quiz completions (1 hr)",
        "Meeting with OM & Marshall: Version control demo (15 min)",
        "Laboratory: alignment, measuring tools, precision, source setup (2 hrs)",
        "Exam/Assignment time block (2 hrs)",
        "Eat (No research!) (2 hrs)",
        "Do NOT do any research today"
    ],
    "Tuesday": [
        "Statistical Mechanics: Start studying (1.5 hrs)",
        "Stat Mech: Redo current chapter homework (1.5 hrs)",
        "Atomic & Nuclear: Redo Ch. 3 HW at 8 PM (2 hrs)",
        "Finalize English homework (due by 5 PM) (2 hrs)",
        "Confirm & finalize data collection plan for source (1 hr)",
        "Ask Berhane if it’s rigorous enough (15 min)",
        "Modern Physics: Homework planning and start submission (1 hr)"
    ],
    "Wednesday": [
        "Continue Statistical Mechanics studying (2 hrs)",
        "Obtain MA 442 notes from someone (15 min)",
        "Understand ½ of MA 442 Study Guide (1.5 hrs)",
        "Modern Physics: Finalize homework for Ch. 1–4 (2 hrs)",
        "Laboratory catch-up or polish if needed (1 hr)",
        "Any English adjustments from poster (optional polish, 30 min)"
    ],
    "Thursday": [
        "Complete 40 practice problems (across courses) (4 hrs total across day)",
        "Final polish of poster or review for class prep (1 hr)",
        "Begin final review sweep of remaining open items (2 hrs)",
        "Free time buffer or catch-up (optional)"
    ],
    "Friday": [
        "Send Leah the Modern Physics HW problems (15 min)",
        "Modern Physics: Make sure all HWs are submitted by EOD (1 hr)",
        "Follow up with Hughes on PS 405 examination (15 min)"
    ]
}

# Write to document
for day, tasks in todo_by_day.items():
    doc.add_heading(day, level=2)
    for task in tasks:
        doc.add_paragraph(f"• {task}", style='List Bullet')

# Save the document
doc.save("Organized_ToDo_List_April8_12.docx")
