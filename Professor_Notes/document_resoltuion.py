from docx import Document

def create_spectrometer_resolution_doc(filename="spectrometer_resolution_notes.docx"):
    doc = Document()
    
    # Title
    doc.add_heading("Estimating Spectrometer Resolution", level=1)

    # Intro
    doc.add_paragraph("To estimate the spectral resolution (R) of a spectrometer using slit width and wavelength range 'eyeballed' from an image or data, you can use a rough method as outlined below.")

    # Section: What You’ll Need
    doc.add_heading("🔧 What You’ll Need to Estimate:", level=2)
    doc.add_paragraph("1. Slit width (w) — in micrometers (μm) or mm.")
    doc.add_paragraph("2. Wavelength range of a feature (Δλ) — estimated visually, e.g., the Na doublet split.")
    doc.add_paragraph("3. Central wavelength (λ) — the wavelength at which resolution is being estimated, e.g., ~589.3 nm for Na doublet.")

    # Section: Step-by-step Calculation
    doc.add_heading("🧮 Step-by-step Calculation:", level=2)
    
    doc.add_heading("1. Eyeball the line separation", level=3)
    doc.add_paragraph("If you can see that the Na doublet (~589.0 nm and ~589.6 nm) is just barely resolved as two peaks, then your resolution in wavelength is:")
    doc.add_paragraph("Δλ ≈ 0.6 nm", style='Intense Quote')

    doc.add_heading("2. Estimate Spectral Resolution (R)", level=3)
    doc.add_paragraph("R = λ / Δλ", style='Intense Quote')
    doc.add_paragraph("If λ = 589.3 nm and Δλ = 0.6 nm, then:")
    doc.add_paragraph("R ≈ 589.3 / 0.6 ≈ 982", style='Intense Quote')

    # Optional section with slit width
    doc.add_heading("🔍 Optional: Slit Width Consideration", level=2)
    doc.add_paragraph("If you know your slit width (w), you can also use it to estimate instrument-limited resolution:")
    doc.add_paragraph("Δλ_instr ≈ (w × dλ/dx) / M", style='Intense Quote')
    doc.add_paragraph("Where:")
    doc.add_paragraph("- w = slit width in mm")
    doc.add_paragraph("- dλ/dx = dispersion in nm/mm (how much the wavelength changes per mm on the detector)")
    doc.add_paragraph("- M = magnification from slit to detector (often ~1)")

    # Explanation of dλ/dx
    doc.add_heading("📐 Understanding dλ/dx (Dispersion)", level=2)
    doc.add_paragraph("The term dλ/dx refers to the **dispersion** of the spectrometer: how much the wavelength changes per unit distance on the detector (e.g., nanometers per millimeter or nanometers per pixel).")
    doc.add_paragraph("It defines the scale of the spectrum: the larger dλ/dx is, the more 'spread out' the spectrum is on your detector.")
    doc.add_paragraph("You can determine this by calibrating your detector with known spectral lines and measuring how far apart (in mm or pixels) those lines appear.")
    doc.add_paragraph("For example, if two known spectral lines 10 nm apart appear 5 mm apart on your detector, then:")
    doc.add_paragraph("dλ/dx = 10 nm / 5 mm = 2 nm/mm", style='Intense Quote')
    doc.add_paragraph("A smaller dλ/dx means higher resolution per unit length on your detector, but also means the spectrum will occupy less physical space.")

    # Closing note
    doc.add_paragraph("To get a more accurate estimate, you’ll need to know your spectrometer’s dispersion calibration or pixel-to-wavelength scale.")

    # Save the document
    doc.save(filename)
    print(f"Document saved as '{filename}'")

if __name__ == "__main__":
    create_spectrometer_resolution_doc()
