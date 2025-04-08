from docx import Document
from docx.shared import Pt

# Create a new Word document
doc = Document()

# Title
doc.add_heading("FLI ProLine 4022 CCD Imaging System", level=1)

# Section: Technical Specifications
doc.add_heading("Technical Specifications", level=2)
specs = [
    ("Sensor", "Kodak KAI-4022 CCD"),
    ("Resolution", "2048 x 2048 pixels"),
    ("Pixel Size", "7.4 µm x 7.4 µm"),
    ("Imaging Area", "15.15 mm x 15.15 mm"),
    ("Cooling", "Air cooling up to 60°C below ambient; optional liquid cooling up to 80°C below ambient"),
    ("Readout Speeds", "1.5 MHz (low noise), 12 MHz (fast readout ~2 fps)"),
    ("Read Noise", "As low as 5 electrons"),
    ("Dark Current", "0.1 to 1 e⁻/pixel/sec at -30°C (typical)"),
    ("Full Well Capacity", ">40,000 electrons"),
    ("Anti-Blooming", "Overload margin greater than 100x"),
    ("Interface", "USB 2.0"),
    ("Mounting Options", "Standard C-mount; optional Nikon F-mount or Canon EOS mount"),
    ("External Triggering", "Standard")
]

for key, value in specs:
    doc.add_paragraph(f"{key}: {value}", style='List Bullet')

doc.add_paragraph("Source: Finger Lakes Instrumentation [1]")

# Section: Data Acquisition via USB
doc.add_heading("Data Acquisition via USB", level=2)
steps = [
    ("Install Necessary Software",
     "Download and install the FLI Software Installation Kit, which includes USB drivers and ASCOM drivers. This ensures your computer can communicate with the camera. [2]"),
    ("Connect the Camera",
     "Attach the USB cable between the camera and your computer, and connect the 12V power supply. The camera’s fans should activate, confirming it's operational."),
    ("Verify Connection",
     "Once the software and hardware are installed and connected, the camera should be detected. Use FLI software or compatible third-party programs to acquire images. [3]")
]

for title, description in steps:
    doc.add_paragraph(f"{title}:", style='List Number')
    doc.add_paragraph(description, style='BodyText')

# Section: Notes
doc.add_heading("Additional Notes", level=2)
notes = [
    "Ensure that any third-party software used is compatible with FLI cameras.",
    "If issues occur during setup or use, consult the user manual or contact FLI support."
]


# References
doc.add_heading("References", level=2)
references = [
    "[1] https://www.flicamera.com/spec_sheets/archive/ML4022.pdf",
    "[2] https://www.flicamera.com/support/support.php",
    "[3] https://device.report/m/bfc8cdda8e63adaf2bbc24c259a2967d69b6f058ff926016d3e20a17cde4947d"
]

for ref in references:
    doc.add_paragraph(ref, style='BodyText')

# Save the document
file_path = "CCD_Imager_FLI_ProLine_4022_Report.docx"
doc.save(file_path)

file_path
