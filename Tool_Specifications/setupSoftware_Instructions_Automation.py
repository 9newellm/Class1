from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create the document
doc = Document()

# Set document style to Times New Roman, 12pt
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Set heading colors (blue theme)
def set_heading_color(paragraph, rgb_color):
    run = paragraph.runs[0]
    run.font.color.rgb = rgb_color

# Title
title = doc.add_heading('FLI ProLine 4022 Setup Documentation', level=1)
set_heading_color(title, RGBColor(0, 102, 204))

# Section: Overview
overview = doc.add_heading('1. Overview', level=2)
set_heading_color(overview, RGBColor(0, 102, 204))
doc.add_paragraph(
    "This guide provides a step-by-step process to install and configure the Finger Lakes Instrumentation (FLI) ProLine 4022 CCD camera "
    "on a Windows machine. It follows documentation available at the official FLI support page: https://www.flicamera.com/support/support.php",
)

# Section: System Requirements
requirements = doc.add_heading('2. System Requirements', level=2)
set_heading_color(requirements, RGBColor(0, 102, 204))
doc.add_paragraph("• A Windows PC (Windows 7, 8, 10, or 11, 64-bit preferred)\n"
                  "• USB 2.0 or compatible port\n"
                  "• External power supply for the camera\n"
                  "• Internet connection for downloading drivers and software\n")

# Section: Software Installation
install = doc.add_heading('3. Software Installation', level=2)
set_heading_color(install, RGBColor(0, 102, 204))
doc.add_paragraph("Step 1: Uninstall any previous FLI software if installed.\n"
                  "Step 2: Download the FLI Software Installation Kit from the FLI support site.\n"
                  "• 64-bit Windows: https://www.flicamera.com/software/FLI_Software_Installation_Kit_64bit.exe\n"
                  "Step 3: Run the installer and follow the on-screen instructions.\n"
                  "Step 4: After installation, restart the computer if prompted.\n")

# Section: Camera Connection
connect = doc.add_heading('4. Connecting the Camera', level=2)
set_heading_color(connect, RGBColor(0, 102, 204))
doc.add_paragraph("Step 1: Connect the ProLine 4022 camera to your PC using a USB 2.0 cable.\n"
                  "Step 2: Connect the 12V power supply to the camera.\n"
                  "Step 3: Turn on the power source. The camera's fan should spin up.\n"
                  "Step 4: Windows should recognize the device and finish driver installation automatically.\n")

# Section: Camera Operation with FLIGrab
fligrab = doc.add_heading('5. Capturing Images with FLIGrab', level=2)
set_heading_color(fligrab, RGBColor(0, 102, 204))
doc.add_paragraph("Step 1: Download FLIGrab from the FLI support page.\n"
                  "Step 2: Launch FLIGrab after installation.\n"
                  "Step 3: Select your connected camera from the device list.\n"
                  "Step 4: Set exposure time and parameters.\n"
                  "Step 5: Click the capture button to take an image.\n"
                  "Step 6: Save the image from the file menu.\n")

# Section: Optional SDK Installation
sdk = doc.add_heading('6. Optional: SDK for Custom Development', level=2)
set_heading_color(sdk, RGBColor(0, 102, 204))
doc.add_paragraph("If you plan to control the camera using Python, LabVIEW, or MATLAB:\n"
                  "Step 1: Download the open-source FLI SDK from the support site.\n"
                  "Step 2: Follow the included documentation to integrate with your development environment.\n"
                  "• For Python, ensure proper bindings and device permissions are configured.\n")

# Section: References
references = doc.add_heading('7. Reference', level=2)
set_heading_color(references, RGBColor(0, 102, 204))
doc.add_paragraph("All information and resources are available at the official support site:\n"
                  "https://www.flicamera.com/support/support.php")

# Save as .docx file
file_path = "Software_Run_pthon_docs_FLI_ProLine_4022_Setup_Guide.docx"
doc.save(file_path)

file_path

