from docx import Document
from docx.shared import Pt

def generate_lab_report(filename="Spectrometer_Slit_Width_Report.docx"):
    doc = Document()

    # Set Title
    doc.add_heading("Lab Report: Estimating Optimal Slit Width for Na Doublet Resolution", level=1)

    # Author & Date
    doc.add_paragraph("Author: [Your Name Here]")
    doc.add_paragraph("Date: [Insert Date]")

    # Introduction
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph(
        "The purpose of this report is to estimate the optimal slit width required to resolve the sodium D-line doublet "
        "(centered near 589 nm) using a Czerny-Turner type spectrometer. The resolution goal is to distinguish between "
        "the two closely spaced Na lines at approximately 589.0 nm and 589.6 nm. Accurate determination of the slit width "
        "is crucial in balancing spectral resolution with light throughput."
    )

    # Theory
    doc.add_heading("Theoretical Background", level=2)
    doc.add_paragraph(
        "The spectral resolution of a grating-based spectrometer is influenced by the width of the entrance slit and "
        "the system's dispersion (dλ/dx). The dispersion describes how much the wavelength shifts per unit distance on "
        "the detector, typically given in nm/mm or nm/pixel. The narrower the slit, the finer the resolution, but with "
        "reduced intensity."
    )
    doc.add_paragraph(
        "The relationship between the slit width (w), the instrumental resolution (Δλ), and the system's dispersion is:"
    )
    doc.add_paragraph("    w = Δλ / (dλ/dx)", style='Intense Quote')
    doc.add_paragraph(
        "Where:\n"
        "- w is the slit width in mm\n"
        "- Δλ is the desired wavelength resolution\n"
        "- dλ/dx is the dispersion of the spectrometer (nm/mm)"
    )

    # Calculation
    doc.add_heading("Calculation Example", level=2)
    doc.add_paragraph(
        "Assuming the desired resolution is Δλ = 0.6 nm (sufficient to resolve the sodium doublet), and the dispersion "
        "of the spectrometer is approximately 2 nm/mm, we can estimate the ideal slit width as follows:"
    )
    doc.add_paragraph("    w = 0.6 nm / 2 nm/mm = 0.3 mm = 300 μm", style='Intense Quote')
    doc.add_paragraph(
        "Therefore, a slit width of approximately 300 μm is estimated to be suitable for resolving the Na doublet. "
        "For higher resolution, one could use a narrower slit such as 150 μm, at the expense of signal intensity."
    )

    # Determining Dispersion
    doc.add_heading("Determining Dispersion", level=2)
    doc.add_paragraph(
        "If the dispersion is not known, it can be estimated using known spectral features and their separation on the "
        "detector. For example, if two known lines 10 nm apart are found to be 5 mm apart on the detector, then:"
    )
    doc.add_paragraph("    dλ/dx = 10 nm / 5 mm = 2 nm/mm", style='Intense Quote')
    doc.add_paragraph(
        "Alternatively, if using a digital camera or CCD with known pixel size, dispersion can be approximated by:"
    )
    doc.add_paragraph("    dλ/dx = Bandwidth / (pixels × pixel size in mm)", style='Intense Quote')

    # Conclusion
    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph(
        "The estimated slit width required to resolve the sodium doublet using a Czerny-Turner spectrometer is "
        "approximately 300 μm, assuming a dispersion of 2 nm/mm. For enhanced resolution, smaller slits can be used "
        "if light levels allow. Accurate dispersion calibration is essential for optimizing this estimation."
    )

    # Save
    doc.save(filename)
    print(f"Lab report saved as '{filename}'")

if __name__ == "__main__":
    generate_lab_report()
