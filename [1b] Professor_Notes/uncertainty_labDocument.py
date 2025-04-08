from docx import Document
from docx.shared import Pt

def generate_lab_report(filename="Slit_Width_Uncertainty_Report.docx"):
    doc = Document()

    # Set Title
    doc.add_heading("Lab Report: Uncertainty in Slit Width Measurement for Spectrometer Resolution", level=1)

    # Author & Date
    doc.add_paragraph("Author: [Your Name Here]")
    doc.add_paragraph("Date: [Insert Date]")

    # Introduction
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph(
        "In this report, we estimate the uncertainty involved in measuring the slit width for a spectrometer setup. "
        "The slit width is crucial for resolving spectral lines, and in this case, we are focusing on the sodium D-line doublet, "
        "located at approximately 589.0 nm and 589.6 nm. Accurate measurement of the slit width is vital for achieving the desired "
        "spectral resolution, and this report will assess the effect of measurement uncertainty on the ability to resolve this doublet."
    )

    # Theory of Uncertainty
    doc.add_heading("Theory of Uncertainty", level=2)
    doc.add_paragraph(
        "When using a measuring tool with finite resolution, there is an associated uncertainty in the measurement. "
        "For example, a tool with a resolution of 0.5 mm means that any measurement made with the tool could vary by ±0.25 mm, "
        "since the smallest unit the tool can measure is 0.5 mm. This uncertainty is crucial for interpreting experimental results, "
        "especially when precision is required, such as in the case of determining the slit width for resolving closely spaced spectral lines."
    )

    # Given Information
    doc.add_heading("Given Information", level=2)
    doc.add_paragraph(
        "In our case, we are using a measuring tool with a resolution of 0.5 mm to determine the slit width. "
        "We aim to measure the slit width that would allow us to resolve the sodium D-line doublet, which is approximately 589.0 nm and 589.6 nm apart, "
        "with a target resolution of 0.6 nm. The spectrometer dispersion is assumed to be 2 nm/mm, and we have previously estimated "
        "that the required slit width for this resolution is around 0.3 mm."
    )

    # Uncertainty Calculation
    doc.add_heading("Uncertainty Calculation", level=2)
    doc.add_paragraph(
        "The uncertainty in measuring the slit width is determined by the resolution of the tool. Since the tool has a resolution of 0.5 mm, "
        "the uncertainty in the measurement is taken as ± half the smallest unit of measurement. This gives an uncertainty of ±0.25 mm."
    )

    doc.add_paragraph(
        "Now, we calculate the uncertainty in the slit width measurement using the following relationship:"
    )
    doc.add_paragraph(
        "    δw = ± 0.5 mm / 2 = ± 0.25 mm",
        style='Intense Quote'
    )

    doc.add_paragraph(
        "Thus, the uncertainty in the slit width measurement is ±0.25 mm. If we measure the slit width to be 0.3 mm, this means "
        "the actual slit width could vary between 0.05 mm and 0.55 mm. This range represents a significant uncertainty compared to the "
        "target slit width of 0.3 mm."
    )

    # Impact on Spectral Resolution
    doc.add_heading("Impact on Spectral Resolution", level=2)
    doc.add_paragraph(
        "The large uncertainty in slit width measurements has important implications for the spectrometer's ability to resolve spectral lines. "
        "For example, if the slit width is measured as 0.3 mm but can vary by ±0.25 mm, the actual slit width could be between 0.05 mm and 0.55 mm. "
        "This level of uncertainty is large compared to the required precision for resolving the sodium D-line doublet, which is separated by just 0.6 nm."
    )
    doc.add_paragraph(
        "Therefore, achieving the necessary resolution with this measuring tool would be challenging, as the uncertainty in the slit width "
        "measurement exceeds the precision required to resolve the spectral lines. To achieve better precision, a tool with a smaller resolution, "
        "such as a micrometer or a digital caliper, would be needed."
    )

    # Conclusion
    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph(
        "In conclusion, the uncertainty in measuring the slit width with a tool that has a resolution of 0.5 mm is significant enough to "
        "prevent precise determination of the slit width needed to resolve the sodium D-line doublet. With an uncertainty of ±0.25 mm, "
        "the measurement could vary considerably from the target slit width of 0.3 mm, making it unsuitable for high-precision spectroscopic measurements. "
        "A more precise measuring tool is recommended for accurate determination of the slit width in this application."
    )

    # Save the document
    doc.save(filename)
    print(f"Lab report saved as '{filename}'")

if __name__ == "__main__":
    generate_lab_report()
