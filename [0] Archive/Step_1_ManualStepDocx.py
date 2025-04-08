from docx import Document

def create_document():
    # Create a new Document
    doc = Document()
    doc.add_heading('Manual Steps for Enabling .NET Framework and Registering mscoree.dll', 0)

    # Introduction
    doc.add_paragraph(
        "This document outlines the manual procedures to enable the required .NET Framework version "
        "and register the mscoree.dll file, corresponding to the automated tasks performed by the Python script."
    )

    # Step 1: Check Installed .NET Framework Versions
    doc.add_heading('Step 1: Check Installed .NET Framework Versions', level=1)
    doc.add_paragraph(
        "To verify which versions of the .NET Framework are installed on your system:"
    )
    doc.add_paragraph(
        "1. Press `Win + R` to open the Run dialog box."
        "\n2. Type `regedit` and press Enter to open the Registry Editor."
        "\n3. Navigate to the following path:"
        "\n   `HKEY_LOCAL_MACHINE\\SOFTWARE\\Microsoft\\NET Framework Setup\\NDP\\v4\\Full`"
        "\n4. In the right pane, locate the 'Version' entry to see the installed version."
    )

    # Step 2: Enable the Required .NET Framework Version
    doc.add_heading('Step 2: Enable the Required .NET Framework Version', level=1)
    doc.add_paragraph(
        "If the required .NET Framework version is not installed or enabled:"
    )
    doc.add_paragraph(
        "1. Press `Win + R` to open the Run dialog box."
        "\n2. Type `optionalfeatures` and press Enter to open the Windows Features dialog."
        "\n3. In the dialog, locate the desired version of the .NET Framework (e.g., '.NET Framework 4.8')."
        "\n4. Ensure the checkbox next to it is checked."
        "\n5. Click OK and restart your computer if prompted."
    )

    # Step 3: Register mscoree.dll
    doc.add_heading('Step 3: Register mscoree.dll', level=1)
    doc.add_paragraph(
        "To register the mscoree.dll file:"
    )
    doc.add_paragraph(
        "1. Press `Win + X` and select 'Command Prompt (Admin)' or 'Windows PowerShell (Admin)'."
        "\n2. In the command window, type the following command and press Enter:"
        "\n   `regsvr32 mscoree.dll`"
        "\n3. You should receive a message indicating that the registration was successful."
    )

    # Save the document
    doc.save('Manual_Steps_Documentation.docx')
    print("Documentation created successfully.")

if __name__ == "__main__":
    create_document()
