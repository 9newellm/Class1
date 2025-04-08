from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('Analysis of Sodium Spectral Lines: D-Lines (Doublet) and S-Lines', 0)

# Sodium D-Lines (Doublet)
doc.add_heading('1. Sodium D-Lines (Doublet):', level=1)
doc.add_paragraph(
    "Origin: The D-lines result from transitions between the 3p and 3s energy levels "
    "of sodium atoms, influenced by spin-orbit coupling, leading to fine structure splitting."
)
doc.add_paragraph(
    "Expected Wavelengths:\n"
    "- D₂ Line: Approximately 589.5924 nm (5895.924 Å)\n"
    "- D₁ Line: Approximately 588.9950 nm (5889.950 Å)"
)
doc.add_paragraph(
    "Intensity Ratio: The D₁ line at 589.0 nm is typically twice as intense as the D₂ line at 589.6 nm."
)

# Sodium S-Lines
doc.add_heading('2. Sodium S-Lines:', level=1)
doc.add_paragraph(
    "Origin: The S-lines are ultraviolet (UV) emissions resulting from transitions from "
    "higher principal quantum numbers (n > 3) to the 3s state in sodium atoms."
)
doc.add_paragraph(
    "Expected Wavelengths (calculated using the Rydberg formula):\n"
    "- Transition from n=4 to n=3: Wavelength ≈ 15.49 nm\n"
    "- Transition from n=5 to n=3: Wavelength ≈ 10.59 nm\n"
    "- Transition from n=6 to n=3: Wavelength ≈ 9.04 nm\n"
    "- Transition from n=7 to n=3: Wavelength ≈ 8.30 nm\n"
    "- Transition from n=8 to n=3: Wavelength ≈ 7.89 nm\n"
    "- Transition from n=9 to n=3: Wavelength ≈ 7.63 nm\n"
    "- Transition from n=10 to n=3: Wavelength ≈ 7.45 nm"
)
doc.add_paragraph(
    "Experimental Considerations:\n"
    "- Spectrometer Calibration: Ensure that your Czerny-Turner spectrometer is properly calibrated to accurately measure the wavelengths of both D and S lines.\n"
    "- Detection of S-Lines: Given their position in the UV spectrum, detecting sodium S-lines requires a spectrometer equipped with a detector sensitive to UV wavelengths.\n"
    "- Expected Observations: With appropriate experimental conditions and equipment, you should observe the D-lines in the visible spectrum and, if your setup allows, the S-lines in the UV range."
)

# Save the document
doc.save('sodium_spectral_lines_analysis.docx')
