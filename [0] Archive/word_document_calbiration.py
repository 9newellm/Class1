from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('Calibration of Czerny-Turner Spectrometer', 0)

# Add content to the document
doc.add_paragraph("Calibrating a Czerny-Turner spectrometer involves ensuring that the measured wavelengths are accurate "
                  "and correspond to known reference values. Here’s a step-by-step guide on how to calibrate a Czerny-Turner spectrometer:\n")

# Step 1
doc.add_heading('1. Select a Calibration Standard', level=1)
doc.add_paragraph("Choose a known light source with well-defined spectral lines. This can be:\n"
                  "- **Mercury (Hg) lamp**: A mercury vapor lamp emits light at well-known wavelengths (e.g., 365.5 nm, 404.7 nm, 435.8 nm, etc.).\n"
                  "- **Neon (Ne) lamp**: Neon lamps are also commonly used and have defined spectral lines (e.g., 585.2 nm, 640.2 nm).\n"
                  "- **Argon (Ar) or other gas discharge lamps**: These can also be used if their emission lines are known.\n\n"
                  "Make sure that the emission lines of the chosen calibration standard are within the spectrometer’s operational wavelength range.")

# Step 2
doc.add_heading('2. Prepare the Spectrometer', level=1)
doc.add_paragraph("• Set up the Czerny-Turner spectrometer: Ensure that the spectrometer is properly aligned, and the optics are in good condition. "
                  "Adjust the slit width and any other setup variables as required for your experiment.\n"
                  "• Choose a suitable wavelength range: Set your spectrometer to scan the wavelength range that encompasses the lines from the calibration source.")

# Step 3
doc.add_heading('3. Collect Calibration Data', level=1)
doc.add_paragraph("• Turn on the calibration light source (e.g., Hg or Ne lamp).\n"
                  "• Run a scan with the spectrometer and record the intensity vs. wavelength data. This will give you a spectrum with peaks at known wavelengths.\n"
                  "• Identify the peaks in the collected spectrum, and match them with the known wavelengths of the calibration source.")

# Step 4
doc.add_heading('4. Correct for Spectrometer Offset', level=1)
doc.add_paragraph("• **Fit the peaks to the known wavelengths**: The peaks from the collected data should ideally correspond to the known emission lines of the calibration source. "
                  "If there is a deviation, you can apply a **linear correction factor** to adjust the spectrometer's wavelength calibration.\n"
                  "• The formula for this correction is:\n"
                  "  **λ_measured = λ_true + offset**\n"
                  "    where:\n"
                  "    - λ_measured is the wavelength as measured by the spectrometer\n"
                  "    - λ_true is the known or expected wavelength from the calibration source\n"
                  "    - **offset** is the constant difference that needs to be corrected.\n"
                  "• You can calculate the **offset** by comparing the measured wavelength and the true wavelength for several calibration peaks. "
                  "If the discrepancy is consistent, you can use it to adjust all measured wavelengths.")

# Step 5
doc.add_heading('5. Apply Calibration to All Measurements', level=1)
doc.add_paragraph("• Once you have the calibration curve or correction factor (offset), apply it to all future spectral measurements. This ensures that the measured wavelengths are "
                  "accurately aligned with the true wavelengths.")

# Step 6
doc.add_heading('6. Perform Regular Calibration', level=1)
doc.add_paragraph("• It is recommended to calibrate the spectrometer periodically, especially if you notice drift in the measurements or after maintenance activities. "
                  "Calibration should be checked regularly to maintain accuracy over time.")

# Step 7
doc.add_heading('7. Documentation', level=1)
doc.add_paragraph("• Record the calibration procedure, including the light source used, the known emission lines, the measured peaks, and any corrections or offsets applied. "
                  "This will help with reproducibility and troubleshooting in the future.")

# Additional Considerations
doc.add_heading('Additional Considerations', level=1)
doc.add_paragraph("• **Spectral Resolution**: The resolution of the spectrometer can affect the accuracy of the wavelength measurement. Ensure that your spectrometer’s resolution is sufficient "
                  "to distinguish the calibration peaks clearly.\n"
                  "• **Temperature Effects**: The spectrometer's optical components may drift with temperature. It's important to calibrate under the same environmental conditions as your experiment.")

# Example of Calibration Process
doc.add_heading('Example of Calibration Process', level=1)
doc.add_paragraph("Let’s say you are using a **Mercury (Hg) lamp** with known emission lines at **435.8 nm** and **546.1 nm**:\n"
                  "1. Run the spectrometer with the Hg lamp and record the spectrum.\n"
                  "2. Find the measured peak positions, say the first peak is at **436.2 nm** and the second at **547.5 nm**.\n"
                  "3. Calculate the offsets:\n"
                  "   **Offset for 435.8 nm = 436.2 - 435.8 = 0.4 nm**\n"
                  "   **Offset for 546.1 nm = 547.5 - 546.1 = 1.4 nm**\n"
                  "4. Take the average offset and apply it to future measurements to correct for any spectral errors.")

# Save the document
doc_path = "Czerny_Turner_Spectrometer_Calibration_Guide.docx"
doc.save(doc_path)

doc_path
