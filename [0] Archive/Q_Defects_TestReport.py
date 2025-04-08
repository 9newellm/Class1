from docx import Document
import pandas as pd 

def create_test_report(file_name):
    # Create a new document
    doc = Document()
    
    # Add report title
    doc.add_heading('Raman Spectrometer Quantum Defect Test Report', 0)
    
    # Add report details
    doc.add_paragraph('Test Report Identifier: RSP-001')
    doc.add_paragraph('Test Date: [Insert Date]')
    doc.add_paragraph('Test Location: [Insert Location]')
    doc.add_paragraph('Test Engineer: Madison J. Newell')
    doc.add_paragraph('Quality Engineer: Berhane')

    doc.add_heading('1. Test Objective', level=1)
    doc.add_paragraph('To document the performance evaluation of a Raman spectrometer using a 527 nm excitation laser, '
                      'including spectral accuracy, resolution, and efficiency based on the test procedure RSP-001.')
    
    doc.add_heading('2. Test Conditions', level=1)
    doc.add_paragraph('Environmental Conditions: [Temperature, Humidity, Pressure]')
    doc.add_paragraph('Instrument Calibration: Verified using a silicon wafer as a reference sample.')
    doc.add_paragraph('Alignment Checks: Ensured proper positioning of collimating and focusing mirrors.')

    doc.add_heading('3. Data Analysis and Results', level=1)
    
    # Add sample data analysis
    doc.add_paragraph('The quantum defect for each sample has been calculated using the Raman peak positions. '
                      'Below are the results of the analysis:')
    
    doc.add_paragraph('Sample Data:')
    
    # Add the analysis results to the report (this part should reflect the actual data from your analysis)
    doc.add_paragraph('The analysis results are shown below (simplified table):')
    
    # For this example, we're using the analysis results from the Excel file (assuming it's already analyzed)
    analysis_results = pd.read_excel('raman_analysis_results.xlsx')
    
    # Create a table in the document
    table = doc.add_table(rows=1, cols=len(analysis_results.columns))
    table.style = 'Table Grid'
    
    # Add the header row
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(analysis_results.columns):
        hdr_cells[i].text = col
    
    # Add data rows
    for _, row in analysis_results.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    # Save the document
    doc.save('Raman_Spectrometer_Test_Report.docx')

# Create the test report
create_test_report("raman_analysis_results.xlsx")
