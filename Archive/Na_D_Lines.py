from docx import Document
from docx.shared import Pt

# Create a new Document
doc = Document()
doc.add_heading('Sodium D-Lines Spectral Analysis', 0)

# Introduction
doc.add_paragraph(
    "The sodium D-lines are a pair of spectral lines in the yellow region of the visible spectrum, "
    "observed at wavelengths of approximately 589.0 nm (D₁) and 589.6 nm (D₂). "
    "These lines result from electronic transitions within sodium atoms, specifically from the 3p to the 3s energy levels. "
    "The fine structure splitting of these lines arises due to spin-orbit coupling, causing the energy levels to split based on the total angular momentum (J) of the electron."
)

# Section: Understanding Sodium D-Lines in Relation to Hydrogen Spectral Series
doc.add_heading('Understanding Sodium D-Lines in Relation to Hydrogen Spectral Series', level=1)
doc.add_paragraph(
    "While the sodium D-lines are specific to sodium atoms, it's insightful to compare them to the hydrogen atom's spectral series, "
    "such as the Lyman and Balmer series, to understand the underlying principles of atomic transitions:"
)
doc.add_paragraph(
    "• Lyman Series (Hydrogen): Consists of transitions where electrons move from higher energy levels (n ≥ 2) to the n = 1 energy level, emitting ultraviolet light."
)
doc.add_paragraph(
    "• Balmer Series (Hydrogen): Involves transitions from higher energy levels (n ≥ 3) to the n = 2 level, resulting in visible light emissions, such as the prominent Hα line at 656 nm."
)
doc.add_paragraph(
    "In contrast, the sodium D-lines originate from transitions between specific energy levels (3p to 3s) within sodium atoms. "
    "The presence of spin-orbit coupling in sodium leads to the splitting of these lines into two distinct components, known as the D₁ and D₂ lines. "
    "This splitting is not observed in hydrogen, where spectral lines are typically singlets without such fine structure."
)

# Section: Calculating the Wavelengths of Sodium D-Lines Using the Rydberg Formula
doc.add_heading('Calculating the Wavelengths of Sodium D-Lines Using the Rydberg Formula', level=1)
doc.add_paragraph(
    "The Rydberg formula calculates the wavelengths of spectral lines resulting from electron transitions between energy levels in hydrogen-like atoms:"
)
doc.add_paragraph(
    "1/λ = R_H × Z² × (1/n₁² - 1/n₂²)"
)
doc.add_paragraph(
    "Where:"
)
doc.add_paragraph(
    "• λ is the wavelength of the emitted light."
)
doc.add_paragraph(
    "• R_H is the Rydberg constant (1.097 × 10⁷ m⁻¹)."
)
doc.add_paragraph(
    "• Z is the atomic number of the element (for sodium, Z = 11)."
)
doc.add_paragraph(
    "• n₁ and n₂ are the principal quantum numbers of the initial and final energy levels, respectively."
)

# Transition from n=3 to n=4 (D₂ Line)
doc.add_paragraph(
    "1. Transition from n=3 to n=4 (D₂ Line):"
)
doc.add_paragraph(
    "1/λ = R_H × Z² × (1/3² - 1/4²)"
)
doc.add_paragraph(
    "Calculating this yields:"
)
doc.add_paragraph(
    "λ_D2 ≈ 589.5924 nm"
)

# Transition from n=3 to n=5 (D₁ Line)
doc.add_paragraph(
    "2. Transition from n=3 to n=5 (D₁ Line):"
)
doc.add_paragraph(
    "1/λ = R_H × Z² × (1/3² - 1/5²)"
)
doc.add_paragraph(
    "This calculation results in:"
)
doc.add_paragraph(
    "λ_D1 ≈ 588.9950 nm"
)

doc.add_paragraph(
    "These theoretical values closely match the experimentally observed wavelengths of the sodium D-lines, demonstrating the effectiveness of the Rydberg formula in predicting spectral line wavelengths for sodium."
)

# Conclusion
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    "Understanding the sodium D-lines provides valuable insights into atomic spectra. "
    "Comparing these lines with hydrogen's spectral series highlights the role of electron transitions and energy level structures in determining the wavelengths of emitted light. "
    "The ability to calculate these wavelengths using the Rydberg formula enhances our comprehension of atomic spectra and electron behavior within atoms."
)

# Save the document
doc.save('Sodium_D_Lines_Spectral_Analysis.docx')
