from docx import Document

from FOLDER_CONFIG import *

# Create a new Document
doc = Document()
doc.add_heading('Major Fraunhofer Lines', 0)

# Add a table
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

# Set column names as headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Line Designation'
hdr_cells[1].text = 'Associated Element'
hdr_cells[2].text = 'Wavelength (nm)'

# Data for the table
fraunhofer_lines = [
    ('A', 'O₂', 759.37),
    ('B', 'O₂', 686.72),
    ('C', 'Hα (Hydrogen Alpha)', 656.28),
    ('D', 'Sodium D-lines', 589.29),
    ('E', 'Fe', 527.04),
    ('F', 'Hβ (Hydrogen Beta)', 486.13),
    ('G', 'Hγ (Hydrogen Gamma)', 434.05),
    ('H', 'Hδ (Hydrogen Delta)', 410.18),
    ('K', 'Ca²⁺ (Calcium Ion)', 393.37),
]

# Populate the table with data
for line, element, wavelength in fraunhofer_lines:
    row_cells = table.add_row().cells
    row_cells[0].text = line
    row_cells[1].text = element
    row_cells[2].text = str(wavelength)


if NEWCONFIG_STRUCUTURE: 
    # Save the document
    doc.save('Fraunhofer_Lines_Table.docx')
