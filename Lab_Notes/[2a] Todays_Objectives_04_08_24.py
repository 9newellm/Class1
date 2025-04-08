from docx import Document

# Create a new Document
doc = Document()

# Add title
doc.add_heading('To-Do List for Lab (04/08/2025)', 0)

# Add the To-Do list
to_do_list = [
    "1. Align effectively.",
    "2. Include 15.7mm ÷ 10 = 1.57cm for CCD camera depth.",
    "3. Write up alignment process.",
    "4. Image dots for alignment verification.",
    "5. Adjust light source to determine projection.",
    "6. Prepare for potential secondary light source issues.",
    "7. [Pause].",
    "8. Run and capture measurement data (if possible).",
    "9. Adjust slit to calculated resolution size.",
    "10. Verify alignment; check for power/ imaging issues.",
    "11. Capture final data and confirm with Berhane.",
    "12. **Note**: Only 2 more hours left for the lab."
]

# Add each item in the to-do list to the document
for item in to_do_list:
    doc.add_paragraph(item)

# Save the document
doc.save("[2b] Lab_To_Do_List_04_08_2025.docx")

print("Word document created successfully!")
