from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Title
doc.add_paragraph('Tools Required for Telescope System Calibration and Measurement Procedure', style='Heading 1')

# Add some introductory text
doc.add_paragraph(
    'This document lists the tools required for measuring various parameters during the calibration and analysis of the telescope system. '
    'These measurements will be used for spectral resolution analysis and quantum defect calculation.'
)

# Section for Tools
doc.add_paragraph('Tools for Measurement Categories', style='Heading 2')

# Create a table for Tools
table = doc.add_table(rows=1, cols=3)

# Add column headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Measurement Category'
hdr_cells[1].text = 'Tools Required'
hdr_cells[2].text = 'Notes'

# List of categories and tools
tool_data = [
    ("Slit and Optical System", 
     "Micrometer or Vernier Caliper, Optical Ruler/Stage, Optical Bench", 
     "For measuring slit width, distance, and focal length."),
    ("CCD and Detector", 
     "High-resolution microscope, Digital Caliper, Image Processing Software", 
     "For measuring pixel size, resolution, and analyzing SNR."),
    ("Light Source", 
     "Monochromator, Optical Power Meter, Spectrometer", 
     "For measuring wavelength range, intensity profile, and spectral line width."),
    ("Optical Alignment", 
     "Collimator, Laser Pointer, Focusing Stage", 
     "For checking collimation and focus accuracy."),
    ("Spectral Calibration and Measurement", 
     "Spectrometer, Calibration Light Source, Monochromator", 
     "For spectral line separation and calibration."),
    ("Data Analysis", 
     "Image Processing Software (ImageJ, MATLAB, Python)", 
     "For line width analysis, spectral feature resolution, and error analysis."),
    ("Environmental and Stability Measurements", 
     "Digital Thermometer, Hygrometer, Accelerometer", 
     "For monitoring temperature, humidity, and vibration."),
    ("Timing-based Measurements", 
     "Timing Software, Oscilloscope or Timer", 
     "For measuring exposure and integration time."),
]

# Add the data to the table
for row_data in tool_data:
    row = table.add_row().cells
    for i, data in enumerate(row_data):
        row[i].text = data

# Format the table text to Times New Roman and 12 pt size
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# Set table style to Table Grid
table.style = 'Table Grid'

# Set alignment of the table cells to center
for row in table.rows:
    for cell in row.cells:
        # Align each paragraph within the cell to center
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save the document
doc.save("Telescope_System_Tools_Required.docx")

print("Document 'Telescope_System_Tools_Required.docx' has been generated successfully.")
