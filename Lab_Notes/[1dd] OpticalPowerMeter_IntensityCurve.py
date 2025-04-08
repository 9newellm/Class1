from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Title
doc.add_paragraph('Light Source and Optical Power Meter Usage for Spectrometer Calibration', style='Heading 1')

# Introduction Section
doc.add_paragraph(
    'The optical power meter is an essential tool for measuring the optical power (intensity) of light within the optical setup, '
    'especially when working with a spectrometer. This document explains the purpose of the optical power meter, how it should be used, '
    'and where to place it within the system.'
)

# Purpose of Optical Power Meter
doc.add_paragraph('Purpose of the Optical Power Meter', style='Heading 2')
doc.add_paragraph(
    'An optical power meter is used to measure the intensity (optical power) of light in an optical system. '
    'It is critical for ensuring that the correct amount of light reaches the CCD detector and that light is not over or under-exposed, '
    'which would affect the accuracy of spectral measurements, such as resolving spectral lines and determining quantum defects.'
)

# Where and How to Use the Optical Power Meter
doc.add_paragraph('Where and How to Use the Optical Power Meter', style='Heading 2')
doc.add_paragraph(
    '1. **Before the Slit (Source to Slit)**: The optical power meter should be placed before the slit in order to measure the intensity of the light '
    'emitted by the light source. This provides a baseline measurement of how much light is being generated before it passes through the slit.'
)
doc.add_paragraph(
    '2. **At the Exit of the Slit**: The power meter should also be placed after the slit to measure the intensity of light passing through the slit and entering the spectrometer optics. '
    'This helps verify how much light is actually passing through the slit for analysis.'
)
doc.add_paragraph(
    '3. **At the CCD Detector**: Lastly, the optical power meter can be positioned at the CCD detector to measure how much light is being recorded by the detector. '
    'This allows for the optimization of exposure and intensity settings to ensure accurate measurements.'
)

# What You Would Be Measuring
doc.add_paragraph('What You Would Be Measuring', style='Heading 2')
doc.add_paragraph(
    'The optical power meter measures the intensity or optical power of light. In this context, it can help with the following:'
)
doc.add_paragraph(
    '- **Light Intensity**: The power meter measures the light intensity, helping you understand how much light is being captured at various points in the system.'
)
doc.add_paragraph(
    '- **Power Distribution**: By measuring the light power at different points (before the slit, after the slit, and at the detector), the optical power meter helps identify how much light '
    'is lost or modified as it travels through the system.'
)
doc.add_paragraph(
    '- **Spectral Characteristics**: Some optical power meters come with wavelength-dependent features, meaning they can also help you measure how the power is distributed across different wavelengths, '
    'providing additional calibration insights.'
)

# Is the Optical Power Meter the Final Light Source?
doc.add_paragraph('Is the Optical Power Meter the Final Light Source?', style='Heading 2')
doc.add_paragraph(
    'No, the optical power meter is not the final source of light in the setup. The **final light source** refers to the actual light-emitting device you are using in your system, '
    'such as a monochromatic light source, a tungsten-halogen lamp, or a laser.'
)
doc.add_paragraph(
    'The **optical power meter** works alongside the light source to **measure and calibrate** the intensity of light that is emitted and passes through the optical system. '
    'It provides essential data for understanding how much light is reaching different parts of the system (slit, spectrometer, detector) and helps ensure that the measurements are accurate.'
)

# Practical Example
doc.add_paragraph('Practical Example', style='Heading 2')
doc.add_paragraph(
    '1. **Light Source**: A tungsten-halogen lamp could be used as the light source, emitting a broad spectrum of visible light.'
)
doc.add_paragraph(
    '2. **Positioning the Optical Power Meter**: The power meter would first be placed before the slit to measure the intensity of the light from the lamp. The light intensity would be recorded.'
)
doc.add_paragraph(
    '3. **After the Slit**: The power meter would then be placed after the slit to measure the light passing through. This helps verify how much light is allowed through by the slit.'
)
doc.add_paragraph(
    '4. **At the CCD Detector**: Finally, the optical power meter would be placed at the CCD detector to measure the intensity of light reaching the detector.'
)
doc.add_paragraph(
    '5. **Spectral Characteristics**: If you have a monochromator or another spectral tool, the power meter can also help measure the power distribution across different wavelengths to ensure that your measurements are properly calibrated.'
)

# Conclusion Section
doc.add_paragraph('Conclusion', style='Heading 2')
doc.add_paragraph(
    'In summary, the optical power meter is an invaluable tool in spectral measurement setups. It helps ensure that the correct amount of light is used in the spectrometer, '
    'allowing for precise calibration, accurate spectral analysis, and the proper resolution of quantum defects.'
)

# Formatting: Apply Times New Roman font and set the font size
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# Save the document
doc.save("Optical_Power_Meter_Writeup.docx")

print("Document 'Optical_Power_Meter_Writeup.docx' has been generated successfully.")
