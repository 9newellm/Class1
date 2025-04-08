from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH





from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Title
doc.add_paragraph('Tools Required for Telescope System Calibration and Measurement Procedure', style='Heading 1')

# Add some introductory text
doc.add_paragraph(
    'This document lists the tools required for measuring various parameters during the calibration and analysis of the telescope system. '
    'These measurements will be used for spectral resolution analysis and quantum defect calculation.'
)

# Section for Tools
doc.add_paragraph('Tools for Measurement Categories', style='Heading 2')

# Create a table for Tools
table = doc.add_table(rows=1, cols=3)

# Add column headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Measurement Category'
hdr_cells[1].text = 'Tools Required'
hdr_cells[2].text = 'Notes'

# List of categories and tools
tool_data = [
    ("Slit and Optical System", 
     "Micrometer or Vernier Caliper, Optical Ruler/Stage, Optical Bench", 
     "For measuring slit width, distance, and focal length."),
    ("CCD and Detector", 
     "High-resolution microscope, Digital Caliper, Image Processing Software", 
     "For measuring pixel size, resolution, and analyzing SNR."),
    ("Light Source", 
     "Monochromator, Optical Power Meter, Spectrometer", 
     "For measuring wavelength range, intensity profile, and spectral line width."),
    ("Optical Alignment", 
     "Collimator, Laser Pointer, Focusing Stage", 
     "For checking collimation and focus accuracy."),
    ("Spectral Calibration and Measurement", 
     "Spectrometer, Calibration Light Source, Monochromator", 
     "For spectral line separation and calibration."),
    ("Data Analysis", 
     "Image Processing Software (ImageJ, MATLAB, Python)", 
     "For line width analysis, spectral feature resolution, and error analysis."),
    ("Environmental and Stability Measurements", 
     "Digital Thermometer, Hygrometer, Accelerometer", 
     "For monitoring temperature, humidity, and vibration."),
    ("Timing-based Measurements", 
     "Timing Software, Oscilloscope or Timer", 
     "For measuring exposure and integration time."),
]

# Add the data to the table
for row_data in tool_data:
    row = table.add_row().cells
    for i, data in enumerate(row_data):
        row[i].text = data

# Format the table text to Times New Roman and 12 pt size
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# Set table style to Table Grid
table.style = 'Table Grid'

# Set alignment of the table cells to center
for row in table.rows:
    for cell in row.cells:
        cell.text_frame.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save the document
doc.save("Telescope_System_Tools_Required.docx")

print("Document 'Telescope_System_Tools_Required.docx' has been generated successfully.")



"""" Break Point """

# Create a new Document
doc = Document()

# Title
doc.add_paragraph('Test Procedure for Telescope System Calibration and Spectral Resolution Analysis', style='Heading 1')

# Basic Information
doc.add_paragraph("Test Procedure ID: [Your Unique Identifier]")
doc.add_paragraph("Test Engineer: [Your Name]")
doc.add_paragraph("Date: 04/08/2024")
doc.add_paragraph("Professor: [Professor's Name]")
doc.add_paragraph("Course/Project Name: [Course or Project Name]")

# Objective
doc.add_paragraph("1. Objective:")
doc.add_paragraph("This test procedure documents the work performed on the telescope system for the purpose of calibrating "
                  "the optical alignment, improving spectral resolution, and conducting a detailed analysis of quantum defects. "
                  "The goal is to achieve a spectral resolution of 0.5 nm and ensure that the system is properly aligned to analyze "
                  "spectral features, using both a 50 mm slit and a CA-doublet light source.")

# Test Setup and Components
doc.add_paragraph("2. Test Setup and Components:")
doc.add_paragraph("2.1. Telescope System")
doc.add_paragraph("- Optical System: Telescope with adjustable focal length.")
doc.add_paragraph("- Slit Size: 50 mm.")
doc.add_paragraph("- Detector: Flat CCD detector.")
doc.add_paragraph("- Light Source (Initial): White light source.")
doc.add_paragraph("- Light Source (Final): CA-doublet light source.")

doc.add_paragraph("2.2. Test Equipment")
doc.add_paragraph("- Telescope Focal Length: [Enter focal length of your telescope, e.g., 50 mm].")
doc.add_paragraph("- CCD Specifications: [Enter CCD specifications such as pixel size, resolution, etc.].")
doc.add_paragraph("- Slit Tool: [Include specifications, e.g., width, material, etc.].")
doc.add_paragraph("- Optical Components: [List lenses, diffraction gratings, or any other components used in the optical path].")

doc.add_paragraph("2.3. Environmental Conditions")
doc.add_paragraph("- Test Environment: [Enter details, e.g., room temperature, controlled lighting conditions, or any factors that might affect measurements].")
doc.add_paragraph("- External Interference Considerations: [List any factors like ambient light, vibrations, etc., that may impact the results].")

# Procedure Overview
doc.add_paragraph("3. Procedure Overview:")

doc.add_paragraph("3.1. Initial Calibration:")
doc.add_paragraph("Step 1: Set up the telescope system with the 50 mm slit and CCD detector.")
doc.add_paragraph("Step 2: Adjust the telescope to ensure the image is projected clearly onto the CCD, ensuring that the discrete dots of spectral features are visible.")
doc.add_paragraph("Step 3: Record the initial images from the white light source to verify that the telescope system can resolve individual spectral lines.")
doc.add_paragraph("Expected Outcome: Clear and resolved spectral lines on the CCD, showing discrete dots.")

# You can continue adding all the sections following the same pattern...

# Final Notes
doc.add_paragraph("9. Appendix (Optional):")
doc.add_paragraph("A1. CCD Images: [Insert before and after images of spectral lines captured during the test]")
doc.add_paragraph("A2. Calibration Test Data: [Insert the data collected during calibration tests, such as initial spectral line measurements and focusing data]")
doc.add_paragraph("A3. Initial Realignment Images: [Insert images showing the re-alignment process before and after adjustments]")
doc.add_paragraph("A4. Code Utilized: [Insert the code you used for data acquisition, processing, and analysis of the spectral lines]")

# Test Engineer's Signature
doc.add_paragraph("Test Engineer’s Signature:")
doc.add_paragraph("Signature: ________________________")
doc.add_paragraph("Date: _____________________________")

# Formatting the text to Times New Roman and 12 pt size
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# Set the alignment of the document text to justified
for paragraph in doc.paragraphs:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Save the document
doc.save("Telescope_Test_Procedure.docx")


from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Title
doc.add_paragraph('Telescope System Calibration and Measurement Procedure', style='Heading 1')

# Add some introductory text
doc.add_paragraph('This document contains the measurement data table for the telescope system calibration and spectral resolution analysis. '
                  'Fill in the measured values and add comments as necessary.')

# Table for measurements
table = doc.add_table(rows=1, cols=5)

# Add column headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Measurement Category'
hdr_cells[1].text = 'Measurement Parameter'
hdr_cells[2].text = 'Value'
hdr_cells[3].text = 'Unit'
hdr_cells[4].text = 'Notes/Comments'

# Add rows for the data
measurements_data = [
    ("Slit and Optical System", "Slit Width", "[Enter Value]", "mm", "[Enter Notes]"),
    ("Slit and Optical System", "Slit-to-CCD Distance", "[Enter Value]", "mm", "[Enter Notes]"),
    ("Slit and Optical System", "Focal Length", "[Enter Value]", "mm", "[Enter Notes]"),
    ("Slit and Optical System", "Field of View (FoV)", "[Enter Value]", "degrees", "[Enter Notes]"),
    ("CCD and Detector", "Pixel Size", "[Enter Value]", "µm", "[Enter Notes]"),
    ("CCD and Detector", "Pixel Resolution", "[Enter Value]", "pixels", "[Enter Notes]"),
    ("CCD and Detector", "Signal-to-Noise Ratio (SNR)", "[Enter Value]", "dB", "[Enter Notes]"),
    ("CCD and Detector", "Quantum Efficiency", "[Enter Value]", "%", "[Enter Notes]"),
    ("Light Source", "Wavelength Range", "[Enter Value]", "nm", "[Enter Notes]"),
    ("Light Source", "Intensity Profile", "[Enter Value]", "[Enter Units]", "[Enter Notes]"),
    ("Light Source", "Spectral Line Width", "[Enter Value]", "nm", "[Enter Notes]"),
    ("Light Source", "Stability", "[Enter Value]", "[Enter Units]", "[Enter Notes]"),
    ("Optical Alignment", "Collimation Accuracy", "[Enter Value]", "[Enter Units]", "[Enter Notes]"),
    ("Optical Alignment", "Focus Accuracy", "[Enter Value]", "[Enter Units]", "[Enter Notes]"),
    ("Optical Alignment", "Angle of Incidence", "[Enter Value]", "degrees", "[Enter Notes]"),
    ("Optical Alignment", "Optical Path Lengths", "[Enter Value]", "mm", "[Enter Notes]"),
    ("Spectral Calibration", "Spectral Line Separation", "[Enter Value]", "nm", "[Enter Notes]"),
    ("Spectral Calibration", "Spectral Resolution", "[Enter Value]", "nm", "[Enter Notes]"),
    ("Spectral Calibration", "Wavelength Calibration", "[Enter Value]", "nm", "[Enter Notes]"),
    ("Data Analysis", "Line Width (FWHM)", "[Enter Value]", "nm", "[Enter Notes]"),
    ("Data Analysis", "Pixel Resolution on Spectral Features", "[Enter Value]", "pixels", "[Enter Notes]"),
    ("Data Analysis", "Error Analysis", "[Enter Value]", "[Enter Units]", "[Enter Notes]"),
    ("Environmental Stability", "Temperature", "[Enter Value]", "°C", "[Enter Notes]"),
    ("Environmental Stability", "Humidity", "[Enter Value]", "%", "[Enter Notes]"),
    ("Environmental Stability", "Vibration Levels", "[Enter Value]", "[Enter Units]", "[Enter Notes]"),
    ("Environmental Stability", "Light Fluctuations", "[Enter Value]", "[Enter Units]", "[Enter Notes]"),
    ("Timing", "Exposure Time", "[Enter Value]", "seconds", "[Enter Notes]"),
    ("Timing", "Integration Time", "[Enter Value]", "seconds", "[Enter Notes]"),
]

# Add data to the table
for row_data in measurements_data:
    row = table.add_row().cells
    for i, data in enumerate(row_data):
        row[i].text = data

# Format the text to Times New Roman and 12 pt size
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# Set the table formatting
table.style = 'Table Grid'

# Set alignment of the table cells to center
for row in table.rows:
    for cell in row.cells:
        cell.text_frame.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save the document
doc.save("Telescope_Test_Procedure_Measurements.docx")

print("Document 'Telescope_Test_Procedure_Measurements.docx' has been generated successfully.")
