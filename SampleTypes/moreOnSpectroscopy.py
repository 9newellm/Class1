from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('Understanding Atomic Spectroscopy: j, l, s Quantum Numbers and Spectroscopic Notation', 0)

# Add an introduction
doc.add_paragraph(
    "This document explains the relationship between the total angular momentum quantum number (j), "
    "orbital angular momentum quantum number (l), and spin quantum number (s). It also describes how to "
    "extract energy levels and calculate wavelengths of atomic transitions from spectroscopic notation."
)

# Add a section for Relationship Between j, l, and s
doc.add_heading('Relationship Between j, l, and s', level=1)
doc.add_paragraph(
    "In quantum mechanics, the total angular momentum quantum number **j** represents the combined angular momentum resulting from both the orbital angular momentum (**l**) and the intrinsic spin angular momentum (**s**) of an electron. "
    "The relationship between **j**, **l**, and **s** is given by:\n"
    "\n"
    "    j = l ± s\n"
    "\n"
    "This means that **j** can take values derived from adding or subtracting **l** and **s**. For an electron, the spin quantum number **s** is always 1/2. Therefore, for a given orbital angular momentum quantum number **l**, the possible values of **j** are:\n"
    "\n"
    "    j = l + 1/2\n"
    "    j = l - 1/2  (only if l ≥ 1)\n"
    "\n"
    "These possible values of **j** are crucial in determining the energy levels and spectral characteristics of atoms, especially when considering fine structure splitting due to spin-orbit coupling."
)

# Add a section for Extracting Energy Levels and Calculating Wavelengths from Spectroscopic Notation
doc.add_heading('Extracting Energy Levels and Calculating Wavelengths from Spectroscopic Notation', level=1)
doc.add_paragraph(
    "To determine the energy levels and corresponding wavelengths of atomic transitions from spectroscopic notation, follow these steps:\n"
    "\n"
    "1. **Understand Spectroscopic Notation:**\n"
    "   Spectroscopic notation combines the following elements:\n"
    "   - **Term Symbol:** Denoted as (2S+1)L_J, where:\n"
    "     - **(2S+1):** Multiplicity, representing the number of possible orientations of the total spin angular momentum.\n"
    "     - **L:** Total orbital angular momentum, represented by letters (S, P, D, F, ...) corresponding to L = 0, 1, 2, 3, ...\n"
    "     - **J:** Total angular momentum, combining both spin and orbital angular momenta.\n"
    "\n"
    "2. **Determine Quantum Numbers:**\n"
    "   From the term symbol (2S+1)L_J, extract:\n"
    "   - **Spin Quantum Number (S):** Calculate using S = (J - L)/2.\n"
    "   - **Orbital Angular Momentum Quantum Number (L):** Identify from the letter (S, P, D, F, ...) corresponding to L = 0, 1, 2, 3, ...\n"
    "   - **Total Angular Momentum Quantum Number (J):** Given directly in the term symbol.\n"
    "\n"
    "3. **Calculate Energy Levels:**\n"
    "   For a hydrogen-like atom, the energy levels depend solely on the principal quantum number n, given by:\n"
    "\n"
    "       E_n = -13.6 eV / n^2\n"
    "\n"
    "   For multi-electron atoms, energy levels are influenced by electron-electron interactions, and the spectroscopic term provides information about the state of the electron configuration.\n"
    "\n"
    "4. **Determine Transition Wavelengths:**\n"
    "   The wavelength λ of a photon emitted or absorbed during a transition between energy levels E_i and E_f is:\n"
    "\n"
    "       λ = hc / (E_i - E_f)\n"
    "\n"
    "   Where:\n"
    "   - **h:** Planck's constant (6.626 x 10^-34 J·s).\n"
    "   - **c:** Speed of light (3.00 x 10^8 m/s).\n"
    "   - **E_i and E_f:** Initial and final energy levels, respectively.\n"
    "\n"
    "By analyzing spectroscopic notation and applying the above principles, you can extract information about energy levels and calculate the wavelengths of corresponding transitions."
)

# Save the document
doc.save('Atomic_Spectroscopy_Guide.docx')
