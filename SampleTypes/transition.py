# from docx import Document

# # Create a new Document
# doc = Document()
# doc.add_heading('Sodium Spectral Lines: D-Lines and S-Lines', 0)

# # Add an introduction
# doc.add_paragraph(
#     "This document provides a summary of the sodium D-lines and S-lines, "
#     "including their origins, associated transitions, and wavelengths."
# )

# # Add a table for Sodium D-Lines
# doc.add_heading('Sodium D-Lines', level=1)
# table_d = doc.add_table(rows=1, cols=3)
# table_d.style = 'Table Grid'

# # Set column names as headers
# hdr_cells_d = table_d.rows[0].cells
# hdr_cells_d[0].text = 'Line Designation'
# hdr_cells_d[1].text = 'Transition'
# hdr_cells_d[2].text = 'Wavelength (nm)'

# # Data for Sodium D-Lines
# d_lines = [
#     ('D₁', '3p₃/₂ → 3s₁/₂', 589.00),
#     ('D₂', '3p₁/₂ → 3s₁/₂', 589.6)
# ]

# # Populate the table with data
# for line, transition, wavelength in d_lines:
#     row_cells = table_d.add_row().cells
#     row_cells[0].text = line
#     row_cells[1].text = transition
#     row_cells[2].text = str(wavelength)

# # Add a table for Sodium S-Lines
# doc.add_heading('Sodium S-Lines', level=1)
# table_s = doc.add_table(rows=1, cols=3)
# table_s.style = 'Table Grid'

# # Set column names as headers
# hdr_cells_s = table_s.rows[0].cells
# hdr_cells_s[0].text = 'Line Designation'
# hdr_cells_s[1].text = 'Transition'
# hdr_cells_s[2].text = 'Wavelength (nm)'

# # Data for Sodium S-Lines
# s_lines = [
#     ('S₁', '3p₁/₂ → 4s₁/₂', 568.82),
#     ('S₂', '3p₃/₂ → 4s₁/₂', 568.82)
# ]

# # Populate the table with data
# for line, transition, wavelength in s_lines:
#     row_cells = table_s.add_row().cells
#     row_cells[0].text = line
#     row_cells[1].text = transition
#     row_cells[2].text = str(wavelength)

# # Save the document
# doc.save('Sodium_Spectral_Lines.docx')

from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('Sodium Spectral Lines: D-Lines and S-Lines', 0)

# Add an introduction
doc.add_paragraph(
    "This document provides a detailed overview of the sodium D-lines and S-lines, "
    "including their origins, associated transitions, wavelengths, and the quantum numbers involved."
)

# Add a section for Sodium D-Lines
doc.add_heading('Sodium D-Lines', level=1)

# Add a table for Sodium D-Lines
table_d = doc.add_table(rows=1, cols=4)
table_d.style = 'Table Grid'

# Set column names as headers
hdr_cells_d = table_d.rows[0].cells
hdr_cells_d[0].text = 'Line Designation'
hdr_cells_d[1].text = 'Transition'
hdr_cells_d[2].text = 'Wavelength (nm)'
hdr_cells_d[3].text = 'Quantum Numbers (n, l, s, j)'

# Data for Sodium D-Lines
d_lines = [
    ('D₂', '3p₃/₂ → 3s₁/₂', 589.0, 'Initial: n=3, l=1 (p), s=1/2, j=3/2; Final: n=3, l=0 (s), s=1/2, j=1/2'),
    ('D₁', '3p₁/₂ → 3s₁/₂', 589.6, 'Initial: n=3, l=1 (p), s=1/2, j=1/2; Final: n=3, l=0 (s), s=1/2, j=1/2')
]

# Populate the table with data
for line, transition, wavelength, quantum_numbers in d_lines:
    row_cells = table_d.add_row().cells
    row_cells[0].text = line
    row_cells[1].text = transition
    row_cells[2].text = str(wavelength)
    row_cells[3].text = quantum_numbers

# Add a section for Sodium S-Lines
doc.add_heading('Sodium S-Lines', level=1)

# Add a table for Sodium S-Lines
table_s = doc.add_table(rows=1, cols=4)
table_s.style = 'Table Grid'

# Set column names as headers
hdr_cells_s = table_s.rows[0].cells
hdr_cells_s[0].text = 'Line Designation'
hdr_cells_s[1].text = 'Transition'
hdr_cells_s[2].text = 'Wavelength (nm)'
hdr_cells_s[3].text = 'Quantum Numbers (n, l, s, j)'

# Data for Sodium S-Lines
s_lines = [
    ('S₁', '3p₁/₂ → 4s₁/₂', 568.82, 'Initial: n=3, l=1 (p), s=1/2, j=1/2; Final: n=4, l=0 (s), s=1/2, j=1/2'),
    ('S₂', '3p₃/₂ → 4s₁/₂', 568.82, 'Initial: n=3, l=1 (p), s=1/2, j=3/2; Final: n=4, l=0 (s), s=1/2, j=1/2')
]

# Populate the table with data
for line, transition, wavelength, quantum_numbers in s_lines:
    row_cells = table_s.add_row().cells
    row_cells[0].text = line
    row_cells[1].text = transition
    row_cells[2].text = str(wavelength)
    row_cells[3].text = quantum_numbers

# Add a section for Explanation of Wavelength Differences
doc.add_heading('Explanation of Wavelength Differences', level=1)
doc.add_paragraph(
    "The sodium D-lines exhibit a slight difference in wavelengths due to fine structure splitting, "
    "which results from spin-orbit coupling. In the 3p state, the interaction between the electron's spin "
    "and its orbital motion causes the energy levels to split into two levels:"
)
doc.add_paragraph(
    "- **3p₃/₂ (j = 3/2):** Higher energy state.\n"
    "- **3p₁/₂ (j = 1/2):** Lower energy state."
)
doc.add_paragraph(
    "The transition from the 3p₃/₂ state to the 3s₁/₂ state emits the D₂ line, while the transition from "
    "the 3p₁/₂ state to the 3s₁/₂ state emits the D1 line. The energy difference between these levels corresponds "
    "to the observed difference in wavelengths, approximately 0.6 nm."
)

# Add a section for Calculating Wavelength Differences
doc.add_heading('Calculating Wavelength Differences', level=1)
doc.add_paragraph(
    "The energy difference (ΔE) between two levels is related to the wavelength (λ) of the emitted photon by the equation:"
)
doc.add_paragraph(
    "ΔE = hc / λ"
)
doc.add_paragraph(
    "Rearranging to solve for λ:"
)
doc.add_paragraph(
    "λ = hc / ΔE"
)
doc.add_paragraph(
    "Given that the energy difference between the 3p₃/₂ and 3p₁/₂ states is approximately 0.0021 eV, "
    "we can calculate the corresponding wavelength difference. Using the known values of Planck's constant (h) "
    "and the speed of light (c), this calculation yields a wavelength difference of approximately 0.6 nm, "
    "which matches the observed difference between the D₂ and D1 lines."
)

# Save the document
doc.save('Sodium_Spectral_Lines_Quantum_Analysis.docx')

